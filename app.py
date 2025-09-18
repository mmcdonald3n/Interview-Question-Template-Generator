# app.py — Neogen Interview Questions Generator (Streamlit)
# --------------------------------------------------------
# What this app does
# - Upload a Job Description (TXT / DOCX / PDF / MD) OR paste JD text
# - Generates interview questions in Neogen house style, aligned to the "Neogen 1st Interview Guide" structure
# - Legal guardrails for US + EU
# - Preflight JD "Compliance Hints" (flags risky wording like “young”, “recent graduate”, etc.)
# - Download as DOCX or Markdown
#
# Setup (local)
#   pip install -r requirements.txt
#   export OPENAI_API_KEY=your_key_here   # or set in Streamlit Cloud secrets
#   streamlit run app.py

import os
import io
import re
from datetime import datetime
from typing import List, Tuple, Dict

import streamlit as st

# Optional deps; handle gracefully if missing
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except Exception:
    Document = None  # type: ignore

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None  # type: ignore

# OpenAI (new SDK style)
try:
    from openai import OpenAI
    _client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except Exception:
    _client = None

APP_TITLE = "Neogen Interview Questions Generator"
HOUSE_STYLE_NAME = "Neogen House Style"

st.set_page_config(page_title=APP_TITLE, page_icon="🧠", layout="wide")

# ---------------------------
# Logo utilities
# ---------------------------

def get_logo_path() -> str | None:
    candidates = [
        "neogen-logo-green.webp",
        "neogen_logo_green.webp",
        "neogen-logo-green.png",
        "assets/neogen-logo-green.webp",
        "assets/neogen-logo-green.png",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None

def header():
    col_logo, col_title = st.columns([1, 6], vertical_alignment="center")
    with col_logo:
        lp = get_logo_path()
        if lp:
            st.image(lp, use_container_width=True)
    with col_title:
        st.title(APP_TITLE)
        st.caption("Upload a JD and get structured interview questions in Neogen’s house style with US+EU legal guardrails.")

header()

# ---------------------------
# Text extraction
# ---------------------------

def extract_text_from_upload(upload) -> str:
    if upload is None:
        return ""
    name = upload.name.lower()
    data = upload.read()
    if name.endswith(".txt") or name.endswith(".md"):
        try:
            return data.decode("utf-8")
        except Exception:
            return data.decode("latin-1", errors="ignore")
    if name.endswith(".docx"):
        if Document is None:
            st.error("python-docx not installed. Add `python-docx` to requirements.txt and redeploy.")
            return ""
        file_like = io.BytesIO(data)
        doc = Document(file_like)
        return "\n".join(p.text for p in doc.paragraphs)
    if name.endswith(".pdf"):
        if PdfReader is None:
            st.error("pypdf not installed. Add `pypdf` to requirements.txt and redeploy.")
            return ""
        file_like = io.BytesIO(data)
        reader = PdfReader(file_like)
        text = []
        for page in reader.pages:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(text)
    st.warning("Unsupported file type. Please upload TXT, MD, DOCX, or PDF.")
    return ""

# ---------------------------
# Compliance preflight (heuristic)
# ---------------------------

RISKY_PHRASES: Dict[str, str] = {
    r"\byoung\b": "Avoid age implications → use ‘early-career’ role level only if objectively defined.",
    r"\brecent graduate\b": "Avoid age proxy → specify entry-level skills instead.",
    r"\bable\-?bodied\b": "Avoid disability bias → describe essential functions and reasonable accommodations.",
    r"\bmust be a (us|uk|eu) citizen\b": "Citizenship limits can be discriminatory; ask for right-to-work unless legally required (e.g., export controls).",
    r"\bnative (english|speaker)\b": "Language as a proxy for nationality → specify required proficiency level instead.",
    r"\b(no )?criminal record\b": "Ban-the-box concerns → if relevant, state ‘background check may be required’ compliant with local law.",
    r"\bclean driving record\b": "If driving is essential, state requirement neutrally and per local law.",
    r"\b(no )?pregnan\w*\b": "Pregnancy/family status is protected → remove.",
    r"\bmarried|single|divorced\b": "Family/marital status is protected → remove.",
}

def compliance_findings(text: str) -> List[Tuple[str, str]]:
    findings = []
    lowered = text.lower()
    for pattern, advice in RISKY_PHRASES.items():
        m = re.search(pattern, lowered)
        if m:
            snippet = text[max(0, m.start() - 30): m.end() + 30]
            findings.append((pattern.strip("\\b"), f"{advice}  Snippet: …{snippet}…"))
    return findings

# ---------------------------
# Prompting
# ---------------------------

def build_system_prompt() -> str:
    return f"""
You are an expert TA Partner at Neogen Corporation. Write in the {HOUSE_STYLE_NAME}: clear section headers in **bold**, concise bullet points (•), UK English, and a professional but human tone.
LEGAL COMPLIANCE: Only generate questions that are compliant in both the US and Europe. Do NOT ask about protected characteristics (e.g., age, race, colour, national origin/citizenship, religion, sex/gender, sexual orientation, gender identity, pregnancy/family/marital status, disability/health/genetic information), union affiliation, or other locally protected classes. Avoid salary history questions in jurisdictions that restrict them. If right-to-work is relevant, phrase neutrally (e.g., “Are you legally authorised to work in <region>?”) without probing immigration status.
Keep questions specific, practical, and evidence-based. Provide brief follow-ups and inline (Good:) and (Red flag:) cues where helpful. Tailor everything strictly to the JD.
""".strip()

def build_user_prompt(jd_text: str, seniority: str, region: str, per_section: int, include_legal_footer: bool) -> str:
    legal_footer = "" if not include_legal_footer else (
        "\n\n**Compliance Advisory (for interviewer reference)**\n"
        "• Avoid questions touching protected characteristics or salary history (where restricted).\n"
        "• If a function requires background checks or driving, state this neutrally and per local law.\n"
        "• Focus on essential functions, measurable outcomes, and reasonable accommodations where relevant.\n"
    )

    return f"""
JOB DESCRIPTION (verbatim):
---
{jd_text}
---

Context:
- Seniority: {seniority}
- Region/Market context: {region}
- Aim: Produce a practical, low-jargon interview pack aligned to Neogen’s first-interview template. Keep questions specific to the role, tech stack, stakeholders, and outcomes in the JD.
- Quantity: ~{per_section} questions per major section (use judgement based on JD importance).

Deliverables in this exact structure (use **bold** for headers and bullets for lists). Keep bullets punchy. Include brief follow-ups where useful. Prefer concrete, job-relevant prompts over generic ones.

**Introduction (Script, 1–2 mins)**
• One-paragraph welcome and format overview.

**Background & Experience (5–8 mins)**
• 2–3 tailored prompts that surface the MOST relevant prior work for this role (avoid CV walk-throughs). – Include 1 follow-up under each bullet.

**Motivation for Neogen (2–3 mins)**
• 2–3 prompts testing understanding of Neogen and role fit in this region/market.

**Skills & Qualifications (6–8 mins)**
• {per_section} prompts tied directly to the JD must-haves (tools, methods, regulations, stakeholders). – Add 1 follow-up per bullet.

**Company Knowledge (2–3 mins)**
• 2 targeted prompts on how the candidate would contribute to Neogen’s mission in this function.

**Role-Specific Questions (Core, 10–12 mins)**
• {per_section} deep-dive prompts grounded in the JD’s outcomes. Ask for artefacts/metrics/decision criteria.

**Behavioural (Values & Ways of Working, 6–8 mins)**
• {per_section} STAR-oriented prompts mapped to collaboration, customer focus, integrity, growth mindset. – Add a realistic follow-up per bullet.

**Scenario-Based / Problem-Solving (6–8 mins)**
• {per_section} realistic scenarios using JD context (include data/constraints). – Add quick grading cues as (Good:) and (Red flag:).

**Candidate Questions (2–4 mins)**
• 3 suggested questions a strong candidate might ask (for interviewer awareness).

**Conclusion & Next Steps (Script, 1–2 mins)**
• A short close-out script and immediate next steps.

**Evaluation Rubric (Concise)**
• Define 3–5 criteria with descriptors for 1 (Below), 3 (Meets), 5 (Exceeds).

**Scorecard Template & Notes Page**
• Role | Interviewer | Date.
• Sections & Scores (1–5): Background | Motivation | Skills | Company | Role-Specific | Behavioural | Scenario.
• Overall recommendation (Yes/No + 2–3 bullets rationale).
• Notes: 12 numbered lines for handwritten/typed notes.
{legal_footer}

Formatting notes:
- Use bold section headers exactly as shown.
- Use standard bullets (•). Use short, scannable lines in British English.
- No markdown code fences.
""".strip()

def call_llm(prompt: str, system_prompt: str, model: str = "gpt-4o-mini") -> str:
    if _client is None:
        # Fallback if no OPENAI_API_KEY set (so the UI still works for demo)
        return (
            "**Introduction (Script, 1–2 mins)**\n"
            "• Welcome and interview format overview.\n\n"
            "**Background & Experience (5–8 mins)**\n"
            "• Tell me about one project most similar to this role’s remit. – What was your exact scope? – What changed due to your work?\n\n"
            "**Motivation for Neogen (2–3 mins)**\n"
            "• What draws you to this role and the problems we solve?\n\n"
            "**Skills & Qualifications (6–8 mins)**\n"
            "• Walk me through a recent example demonstrating a core JD must-have. – How did you measure success?\n\n"
            "**Company Knowledge (2–3 mins)**\n"
            "• Where could you contribute in your first 90 days and why?\n\n"
            "**Role-Specific Questions (Core, 10–12 mins)**\n"
            "• Deep dive prompt…\n\n"
            "**Behavioural (Values & Ways of Working, 6–8 mins)**\n"
            "• Tell me about a time you influenced without authority. – What would you do differently next time?\n\n"
            "**Scenario-Based / Problem-Solving (6–8 mins)**\n"
            "• Scenario prompt with (Good:) and (Red flag:) cues.\n\n"
            "**Candidate Questions (2–4 mins)**\n"
            "• 3 example candidate questions…\n\n"
            "**Conclusion & Next Steps (Script, 1–2 mins)**\n"
            "• Thank you + next steps script.\n\n"
            "**Evaluation Rubric (Concise)**\n"
            "• Criteria with 1/3/5 descriptors…\n\n"
            "**Scorecard Template & Notes Page**\n"
            "• Role | Interviewer | Date …\n"
        )
    resp = _client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ],
        temperature=0.4,
        max_tokens=2400,
    )
    return resp.choices[0].message.content or ""

# ---------------------------
# Formatting helpers
# ---------------------------

def to_markdown(text: str) -> str:
    text = re.sub(r"```+", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def to_docx(markdown_like: str) -> bytes:
    if Document is None:
        st.error("python-docx not installed. Add `python-docx` to requirements.txt and redeploy.")
        return b""
    doc = Document()
    # Basic style tweaks
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        style.font.size = Pt(11)
    except Exception:
        pass

    lines = [ln for ln in markdown_like.split("\n")]

    def add_bullet(p_text: str, level: int = 1):
        p = doc.add_paragraph(p_text)
        try:
            if level == 1:
                p.style = doc.styles["List Bullet"]
            else:
                p.style = doc.styles["List Bullet 2"]
        except Exception:
            pass

    for ln in lines:
        s = ln.strip()
        if not s:
            doc.add_paragraph("")
            continue
        if s.startswith("**") and s.endswith("**") and len(s) > 4:
            hdr = s.strip("*")
            run = doc.add_paragraph().add_run(hdr)
            run.bold = True
            continue
        if s.startswith("• "):
            add_bullet(s[2:].strip(), level=1)
            continue
        if s.startswith("– ") or s.startswith("- "):
            add_bullet(s[2:].strip(), level=2)
            continue
        doc.add_paragraph(s)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ---------------------------
# UI
# ---------------------------

with st.sidebar:
    st.subheader("Generation Settings")
    seniority = st.selectbox("Seniority", ["Entry", "Associate", "Mid", "Senior", "Manager", "Director", "Executive"], index=3)
    region = st.selectbox("Region / Market Context", ["USA", "Canada", "UK & Ireland", "EMEA", "LATAM", "APAC", "Global"], index=0)
    per_section = st.slider("Questions per section", 3, 10, 5)
    model = st.selectbox("Model", ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "gpt-4.1"], index=0)
    include_legal_footer = st.checkbox("Include Compliance Advisory at end", value=True)
    st.write(":information_source: Uses your OPENAI_API_KEY from environment.")

    if st.button("Clear inputs", type="secondary"):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()

col1, col2 = st.columns(2)
with col1:
    upload = st.file_uploader("Upload Job Description (TXT, DOCX, PDF, MD)", type=["txt", "docx", "pdf", "md"])
with col2:
    jd_text_area = st.text_area("…or paste Job Description text", height=280)

jd_text = ""
if upload is not None:
    jd_text = extract_text_from_upload(upload)
if not jd_text:
    jd_text = jd_text_area

if jd_text:
    st.success("Job Description received.")

    # Compliance preflight
    findings = compliance_findings(jd_text)
    with st.expander("Compliance Hints (from JD text)"):
        if findings:
            for patt, advice in findings:
                st.warning(f"Potential risk: '{patt}'. {advice}")
        else:
            st.info("No obvious issues detected. Still apply local legal judgment.")

    with st.expander("Preview extracted JD text"):
        st.text(jd_text[:4000] + ("\n…" if len(jd_text) > 4000 else ""))

    gen_col1, gen_col2, gen_col3 = st.columns([1,1,1])
    with gen_col1:
        generate_btn = st.button("Generate Interview Pack", type="primary")
    with gen_col2:
        st.write("")
    with gen_col3:
        st.write("")

    if generate_btn:
        with st.spinner("Generating in Neogen house style…"):
            system_prompt = build_system_prompt()
            user_prompt = build_user_prompt(jd_text, seniority, region, per_section, include_legal_footer)
            output = call_llm(user_prompt, system_prompt, model=model)
            md = to_markdown(output)

        st.subheader("Interview Pack Preview")
        st.markdown(md)

        # Downloads
        docx_bytes = to_docx(md)
        stamp = datetime.now().strftime("%Y%m%d_%H%M")
        base_name = f"Neogen_Interview_Questions_{stamp}"

        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            st.download_button(
                label="Download as DOCX",
                data=docx_bytes,
                file_name=f"{base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with dl_col2:
            st.download_button(
                label="Download as Markdown",
                data=md.encode("utf-8"),
                file_name=f"{base_name}.md",
                mime="text/markdown",
            )
else:
    st.info("Upload or paste a Job Description to begin.")
